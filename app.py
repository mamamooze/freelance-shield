import streamlit as st
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import os
import time
import io

# --- 1. SETUP & CONFIG ---
icon_path = "logo.png"
page_icon = icon_path if os.path.exists(icon_path) else "🛡️"

st.set_page_config(
    page_title="Freelance Shield Pro",
    page_icon=page_icon,
    layout="wide",
    initial_sidebar_state="expanded",
)

# --- 2. CUSTOM CSS ---
st.markdown(
    """
    <style>
        /* BACKGROUND */
        .stApp {
            background-image: linear-gradient(rgba(10, 10, 20, 0.85), rgba(10, 10, 20, 0.90)), 
            url("https://raw.githubusercontent.com/mamamooze/freelance-shield/main/background.png");
            background-size: cover;
            background-attachment: fixed;
        }

        /* TYPOGRAPHY */
        h1 {
            background: -webkit-linear-gradient(45deg, #ffffff, #00d2ff);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            font-family: 'Inter', sans-serif;
            font-weight: 900;
            font-size: 3.5rem;
            letter-spacing: -1px;
            text-shadow: 0 2px 10px rgba(0,210,255,0.3);
        }
        h2, h3 { color: #f8f9fa !important; font-family: 'Inter', sans-serif; }
        
        /* TEXT VISIBILITY */
        .stMarkdown p, .stMarkdown li, label {
            color: #e0e0e0 !important;
            font-size: 1.05rem;
            line-height: 1.6;
        }

        /* CARDS */
        .stInfo {
            background-color: rgba(30, 41, 59, 0.6);
            border: 1px solid #334155;
            transition: transform 0.2s ease, border-color 0.2s;
        }
        .stInfo:hover {
            transform: translateY(-3px);
            border-color: #00d2ff;
            box-shadow: 0 4px 15px rgba(0, 210, 255, 0.1);
        }

        /* INPUT FIELDS */
        .stTextInput input, .stNumberInput input, .stTextArea textarea, .stSelectbox div[data-baseweb="select"] {
            background-color: #1e293b !important;
            color: #ffffff !important;
            border: 1px solid #475569 !important;
            border-radius: 8px;
        }
        
        /* PREVIEW BOX */
        .stTextArea textarea {
            font-family: 'Courier New', monospace !important;
            background-color: #0f172a !important;
            color: #93c5fd !important;
            border: 1px solid #3b82f6 !important;
        }

        /* BUTTONS */
        .stButton>button {
            background: linear-gradient(90deg, #2563eb, #00d2ff);
            color: white;
            font-weight: bold;
            border: none;
            padding: 0.8rem 1.5rem;
            border-radius: 8px;
            width: 100%;
            text-transform: uppercase;
            letter-spacing: 1px;
            box-shadow: 0 4px 15px rgba(0, 210, 255, 0.3);
            transition: all 0.3s;
        }
        .stButton>button:hover {
            transform: translateY(-2px);
            box-shadow: 0 6px 20px rgba(0, 210, 255, 0.5);
        }

        /* SIDEBAR */
        [data-testid="stSidebar"] { background-color: #0f172a; border-right: 1px solid #1e293b; }
        [data-testid="stSidebar"] h1, [data-testid="stSidebar"] h2, [data-testid="stSidebar"] p { color: #cbd5e1 !important; }
        
        /* PILL TABS */
        .stTabs [data-baseweb="tab-list"] {
            gap: 10px;
            background-color: rgba(30, 41, 59, 0.5);
            padding: 10px;
            border-radius: 30px;
            border: 1px solid #334155;
        }
        .stTabs [data-baseweb="tab"] {
            height: auto !important;
            padding: 8px 20px;
            background-color: transparent;
            border-radius: 20px;
            color: #cbd5e1;
            border: none;
            font-weight: 600;
        }
        .stTabs [aria-selected="true"] {
            background-color: #3b82f6;
            color: white !important;
            box-shadow: 0 4px 10px rgba(59, 130, 246, 0.3);
        }
        
        /* WARNING BOX */
        .warning-box {
            background-color: rgba(255, 193, 7, 0.15); 
            border-left: 4px solid #ffc107; 
            padding: 15px; 
            margin-bottom: 15px; 
            border-radius: 4px; 
            color: #ffecb3; 
            font-size: 0.95rem;
        }
        
        /* SUCCESS BOX */
        .success-box {
            background-color: rgba(16, 185, 129, 0.15); 
            border-left: 4px solid #10b981; 
            padding: 15px; 
            margin: 15px 0; 
            border-radius: 4px; 
            color: #d1fae5;
        }
        
        /* LEGAL FOOTER */
        .legal-footer {
            font-size: 0.8rem; 
            color: #64748b; 
            text-align: center; 
            margin-top: 50px; 
            padding-top: 20px; 
            border-top: 1px solid #334155;
        }
    </style>
    """,
    unsafe_allow_html=True,
)

# --- 3. STATE & LOGIC ---
if 'scope_text' not in st.session_state: 
    st.session_state.scope_text = ""

# --- ENHANCED DOCX GENERATOR WITH PROFESSIONAL FORMATTING ---
def create_professional_docx(full_text, annexure_text, provider_name, client_name):
    """Creates a professionally formatted Word document"""
    doc = Document()
    
    # TITLE
    title = doc.add_heading('PROFESSIONAL SERVICE AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(37, 99, 235)
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Date: {datetime.date.today().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # Parse and format main content
    for line in full_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        
        # Section headers (numbered)
        if line and line[0].isdigit() and '.' in line[:3]:
            heading = doc.add_heading(line, level=2)
            heading_run = heading.runs[0]
            heading_run.font.size = Pt(13)
            heading_run.font.color.rgb = RGBColor(30, 41, 59)
        # Signature lines
        elif 'SIGNED BY' in line or 'Signature:' in line or 'Date:' in line:
            sig_para = doc.add_paragraph(line)
            sig_para.runs[0].font.size = Pt(11)
            sig_para.runs[0].bold = True
        # Regular content
        else:
            para = doc.add_paragraph(line)
            para.runs[0].font.size = Pt(11)
    
    # PAGE BREAK
    doc.add_page_break()
    
    # ANNEXURE
    annexure_title = doc.add_heading('ANNEXURE A: SCOPE OF WORK', 1)
    annexure_title.runs[0].font.color.rgb = RGBColor(37, 99, 235)
    
    # Format annexure content
    for line in annexure_text.split('\n'):
        if line.strip():
            para = doc.add_paragraph(line.strip())
            para.runs[0].font.size = Pt(11)
    
    # Add signature section for annexure
    doc.add_paragraph()
    doc.add_paragraph('_' * 60)
    
    sig_section = doc.add_paragraph()
    sig_section.add_run(f'\nProvider Signature: _____________________ Date: __________\n')
    sig_section.add_run(f'Name: {provider_name}\n\n')
    sig_section.add_run(f'Client Signature: _____________________ Date: __________\n')
    sig_section.add_run(f'Name: {client_name}')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- ENHANCED PDF WITH BETTER FORMATTING ---
class EnhancedPDF(FPDF):
    def header(self):
        if self.page_no() == 1:
            self.set_font('Arial', 'B', 16)
            self.cell(0, 10, 'PROFESSIONAL SERVICE AGREEMENT', 0, 1, 'C')
            self.ln(5)
    
    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')
    
    def chapter_title(self, title):
        self.set_font('Arial', 'B', 12)
        self.set_fill_color(230, 230, 230)
        self.cell(0, 8, title, 0, 1, 'L', 1)
        self.ln(2)
    
    def chapter_body(self, body):
        self.set_font('Arial', '', 10)
        self.multi_cell(0, 5, body)
        self.ln()

def create_professional_pdf(full_text, annexure_text, provider_name, client_name):
    """Creates a professionally formatted PDF with proper encoding"""
    pdf = EnhancedPDF()
    pdf.add_page()
    
    # Add logo if exists
    if os.path.exists("logo.png"):
        try:
            pdf.image("logo.png", 10, 8, 25)
            pdf.ln(25)
        except:
            pdf.ln(5)
    
    # Date
    pdf.set_font('Arial', '', 11)
    pdf.cell(0, 8, f"Date: {datetime.date.today().strftime('%B %d, %Y')}", 0, 1, 'C')
    pdf.ln(5)
    
    # Parse and format content
    lines = full_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Replace special characters for PDF compatibility
        line = line.replace('₹', 'Rs. ').replace('—', '-').replace('"', '"').replace('"', '"')
        
        # Section headers
        if line and line[0].isdigit() and '.' in line[:3]:
            pdf.chapter_title(line)
        # Signature lines
        elif 'SIGNED BY' in line:
            pdf.ln(5)
            pdf.set_font('Arial', 'B', 11)
            pdf.cell(0, 8, line, 0, 1)
        else:
            pdf.set_font('Arial', '', 10)
            pdf.multi_cell(0, 5, line)
    
    # ANNEXURE
    pdf.add_page()
    pdf.chapter_title('ANNEXURE A: SCOPE OF WORK')
    
    # Clean and format scope
    clean_scope = annexure_text.replace('₹', 'Rs. ').replace('—', '-')
    pdf.chapter_body(clean_scope)
    
    # Signature section
    pdf.ln(10)
    pdf.set_font('Arial', '', 10)
    pdf.cell(0, 6, '_' * 90, 0, 1)
    pdf.ln(5)
    pdf.cell(0, 6, f'Provider Signature: _____________________ Date: __________', 0, 1)
    pdf.cell(0, 6, f'Name: {provider_name}', 0, 1)
    pdf.ln(5)
    pdf.cell(0, 6, f'Client Signature: _____________________ Date: __________', 0, 1)
    pdf.cell(0, 6, f'Name: {client_name}', 0, 1)
    
    return pdf.output(dest='S')

# --- TEMPLATES ---
scope_templates = {
    "Select a template...": "",
    "✍️ Content Writing": """DELIVERABLE: 4 SEO Blog Articles (1000 words each)
- FORMAT: .docx, Grammarly score >90
- TOPICS: Approved by Client in advance
- DELIVERY: 2 articles/week via email
- REVISIONS: 1 round included per article
- EXCLUSIONS: No image sourcing, keyword research, or posting""",
    
    "🎨 Graphic Design": """DELIVERABLE: Logo (PNG/SVG), Business Card (PDF), Banner
- BRIEF: Colors/Fonts provided by Client
- REVISIONS: 3 feedback rounds included (within 2 days)
- DELIVERY: Final files via Google Drive in 7 days
- EXCLUSIONS: No printing costs or stock image purchase""",
    
    "🖼️ UI/UX & Web Design": """DELIVERABLE: Wireframe + UI Kit (5 Screens)
- FORMAT: Figma/Sketch/XD files
- TIMELINE: Initial draft in 5 days
- REVISIONS: 2 rounds included
- EXCLUSIONS: No coding/development included""",
    
    "💻 Web Development": """DELIVERABLE: 5-Page Responsive Website (WordPress)
- SPECS: Speed score >80, Contact Form, About Page
- DELIVERY: Staging link for review, ZIP files after payment
- REVISIONS: 2 rounds included
- EXCLUSIONS: Domain/Hosting fees and content writing not included""",
    
    "📱 App Development": """DELIVERABLE: Android App MVP (5 Core Features)
- SPECS: Compiles on Android 11+, Source Code included
- TIMELINE: Weekly sprints, 30-day bug fix warranty
- EXCLUSIONS: Google Play Store upload fees not included""",
    
    "🎥 Video Editing": """DELIVERABLE: Edit 2 YouTube Videos (max 8 mins)
- FORMAT: MP4, 1080p, Color Graded
- TIMELINE: Draft within 48 hours of receiving raw files
- REVISIONS: 2 feedback rounds included
- EXCLUSIONS: No captions, thumbnails, or stock footage""",
    
    "📱 Social Media Marketing": """DELIVERABLE: 12 Static Posts + 4 Reels (Monthly)
- FORMAT: PNG (1080px) and MP4 (<60s)
- SCHEDULE: 3 posts/week, approved by 25th of prev month
- REVISIONS: 2 rounds per month included
- EXCLUSIONS: No paid ad management or community replies""",
    
    "📈 SEO & Digital Marketing": """DELIVERABLE: SEO Audit (20 pages) + Keyword Plan
- FORMAT: PDF Report, Excel Sheet
- SPECS: 30 priority keywords, competitor analysis
- REVISIONS: 1 round included
- EXCLUSIONS: On-page implementation and backlinks not included""",
    
    "📧 Virtual Assistance": """DELIVERABLE: Daily Admin Tasks (Email/Calendar)
- REPORTING: Daily Excel report, Inbox cleared
- AVAILABILITY: Mon-Fri, 9am-5pm
- EXCLUSIONS: No calls, travel booking, or personal errands""",
    
    "📸 Photography": """DELIVERABLE: 50 Product Shots (Edited)
- FORMAT: High-res JPEGs, 3000px, White Background
- TIMELINE: Edits delivered in 3 days
- REVISIONS: 1 re-edit round per batch of 10
- EXCLUSIONS: No props, prints, or location booking fees""",
    
    "🗣️ Translation": """DELIVERABLE: Translate 10k words (Eng-Hindi) + 2 Transcripts
- FORMAT: Word/TXT files
- ACCURACY: >98% standard
- REVISIONS: 1 review round included
- EXCLUSIONS: No subtitling or legal localization""",
    
    "🎙️ Voice-Over": """DELIVERABLE: 3 Commercial Voice-overs (30s) + 1 Podcast Edit
- FORMAT: WAV/MP3, Commercial rights included
- SCRIPT: Supplied by Client
- REVISIONS: 1 correction round included
- EXCLUSIONS: No music production or mixing"""
}

def update_scope():
    if st.session_state.template_selector != "Select a template...":
        st.session_state.scope_text = scope_templates[st.session_state.template_selector]

# --- ENHANCED SMART CLAUSES WITH ADDITIONAL LEGAL PROTECTIONS ---
def get_smart_clauses(category, rate):
    clauses = {
        "acceptance": f"Client review within 5 days. Silence = Acceptance. 2 revisions included. Extra changes billed at {rate}/hr.",
        "warranty": "Provided 'as-is'. No post-delivery support unless specified in Annexure A.",
        "ip_rights": "Client owns IP only AFTER full payment. Use before payment is Copyright Infringement.",
        "cancellation": "Cancellation after work starts incurs forfeiture of the Advance Payment.",
        "termination": "Provider may terminate with 7 days written notice if Client breaches payment terms. Client owes payment for work completed till termination date."
    }
    
    if category in ["💻 Web Development", "📱 App Development"]:
        clauses["warranty"] = f"BUG FIX WARRANTY: Provider agrees to fix critical bugs reported within 30 days of delivery at no additional cost, provided: (a) the issue existed at time of delivery, (b) no third-party modifications have been made, (c) hosting environment meets original specifications. Feature changes billed at {rate}/hr."
        clauses["ip_rights"] = "CODE OWNERSHIP: Client receives full source code rights and documentation upon final payment. Provider retains rights to use generic libraries, frameworks, and non-proprietary code components in future projects."
    
    elif category in ["🎨 Graphic Design", "🎥 Video Editing", "🖼️ UI/UX & Web Design", "📸 Photography"]:
        clauses["acceptance"] = "CREATIVE APPROVAL: Client must approve initial style/concept within 3 days. Rejections based solely on 'personal taste' or 'preference' after initial approval will be billed as a new Change Order."
        clauses["ip_rights"] = "SOURCE FILES: Final deliverables (PNG, JPG, MP4, etc.) transfer to Client upon payment. Raw source files (PSD, AI, PrProj, RAW, etc.) remain property of Provider unless specifically purchased as an add-on. Provider may showcase work in portfolio with Client permission."
    
    elif category in ["📱 Social Media Marketing", "📈 SEO & Digital Marketing"]:
        clauses["warranty"] = "NO ROI GUARANTEE: Provider does NOT guarantee specific business outcomes including but not limited to: sales conversions, follower growth, engagement rates, search rankings, or revenue generation. Platform algorithms are external factors beyond Provider control."
        clauses["acceptance"] = "APPROVAL WINDOW: Content must be approved 24 hours prior to scheduled publishing deadlines to allow for timely posting. Late approvals may result in rescheduling."
    
    elif category == "✍️ Content Writing":
        clauses["warranty"] = "ORIGINALITY WARRANTY: Provider warrants that all work is original (not plagiarized). Provider is not liable for legal or factual errors in Client-provided source material."
        clauses["acceptance"] = "EDITORIAL REVIEW: Client has 3 days for factual corrections and clarity issues. Stylistic rewrites or tone changes count as a revision round."
    
    elif category == "🗣️ Translation":
        clauses["warranty"] = "ACCURACY WARRANTY: Provider guarantees >98% accuracy for literal translation. Errors discovered within 7 days of delivery will be corrected free of charge. Provider is not responsible for cultural localization unless specifically contracted."
        clauses["cancellation"] = "KILL FEE: Cancellation after work begins incurs 50% of total fee. Cancellation after draft delivery incurs 100% of total fee."
        clauses["acceptance"] = "REVIEW PERIOD: Client has 5 days to review for accuracy. Technical terminology disputes require Client to provide reference materials."
    
    elif category == "🎙️ Voice-Over":
        clauses["acceptance"] = "CORRECTION POLICY: Includes 1 round for pronunciation errors, pacing issues, or technical defects. Script changes or creative direction changes require a new fee."
        clauses["cancellation"] = "KILL FEE: 50% fee if cancelled after work begins. 100% fee if cancelled after recording session is complete."
    
    return clauses

# --- 4. SIDEBAR ---
with st.sidebar:
    if os.path.exists("logo.png"): 
        st.image("logo.png", width=120)
    else:
        st.markdown("## 🛡️ Freelance Shield")
    
    st.markdown("### 🎯 Founder's Mission")
    st.write("**Hi, I'm a Law Student working to empower Indian freelancers.**")
    st.write("Every year, thousands of independent professionals lose income to late payments, scope creep, and unfair contracts.")
    st.write("**I built this platform so every freelancer—designer, developer, writer—can generate a legally binding, MSME-protected contract in seconds.**")
    
    st.markdown("- 🚀 No more chasing payments")
    st.markdown("- 🛡️ No more ignored clauses")
    st.markdown("- 🏛️ Legal terms trusted by MSMEs")
    
    st.caption("**You deserve to get paid on time, every time.**")
    
    st.markdown("---")
    st.markdown("### 🆘 Need Custom Help?")
    st.write("Complex project? High-value contract? Don't risk it.")
    st.link_button("📞 Hire Me for Review (Rs. 499)", "https://wa.me/YOUR_NUMBER_HERE")
    
    st.markdown("---")
    with st.expander("⚖️ Terms & Privacy Policy"):
        st.markdown("""
        **1. TERMS OF USE**
        - **No Legal Advice:** This tool provides automated templates for informational purposes only. It does not constitute legal advice or create an attorney-client relationship.
        - **Prohibited Use:** You may not use this site for unlawful, fraudulent, or commercial scraping purposes.
        - **Limitation of Liability:** We are not liable for disputes arising from the use of these contracts. For high-value projects (>Rs. 1 lakh), consult a qualified lawyer.
        - **Jurisdiction:** Disputes regarding this website are subject to the courts of Bengaluru, Karnataka, India.

        **2. PRIVACY POLICY**
        - **Stateless Architecture:** We operate on a "Privacy by Design" model. We do NOT store, save, or log any personal data (names, fees, scope) you enter.
        - **No Database:** All generation happens instantly in your active browser session. Once you close the tab, your data is permanently deleted.
        - **No Third Parties:** We do not sell or trade user data because we do not collect it.
        - **Analytics:** We may use anonymized analytics to improve the tool (e.g., number of contracts generated) but never personal information.
        """)

# --- 5. MAIN UI ---
c1, c2 = st.columns([2, 1])
with c1:
    st.markdown("# Stop Chasing Payments.")
    st.markdown('<p style="color: #94a3b8; font-size: 1.3rem; margin-top: -10px;">Generate watertight, MSME-protected contracts in 30 seconds.</p>', unsafe_allow_html=True)
    
    if 'has_greeted' not in st.session_state:
        hour = datetime.datetime.now().hour
        greeting = "Good Morning" if 5 <= hour < 12 else "Good Afternoon" if 12 <= hour < 18 else "Good Evening"
        st.toast(f"👋 {greeting}, Freelancer! Let's get you protected.")
        st.session_state.has_greeted = True

st.markdown("""
<div style="display: flex; gap: 15px; margin-bottom: 20px; flex-wrap: wrap;">
    <span style="background: #1e293b; padding: 8px 15px; border-radius: 5px; color: #94a3b8; font-size: 0.9rem;">🏛️ MSME Act Protected</span>
    <span style="background: #1e293b; padding: 8px 15px; border-radius: 5px; color: #94a3b8; font-size: 0.9rem;">👻 Anti-Ghosting Clause</span>
    <span style="background: #1e293b; padding: 8px 15px; border-radius: 5px; color: #94a3b8; font-size: 0.9rem;">🔒 IP Lock Until Payment</span>
    <span style="background: #1e293b; padding: 8px 15px; border-radius: 5px; color: #94a3b8; font-size: 0.9rem;">⚡ 30-Second Generation</span>
</div>
""", unsafe_allow_html=True)

st.markdown("---")

# --- TRUST INDICATORS ---
st.markdown("""
<div class="success-box">
    <b>✅ What Makes This Different from Generic Templates:</b><br/>
    • Industry-specific IP clauses (your work is protected until payment clears)<br/>
    • MSME Act 2006 compliance (3x interest on late payments - enforceable in Indian courts)<br/>
    • Smart legal logic that adapts to your industry risk (Tech gets bug warranties, Creative gets source file protection)<br/>
    • Anti-ghosting protection (14-day auto-termination if client disappears)
</div>
""", unsafe_allow_html=True)

tab1, tab2, tab3 = st.tabs(["👤 The Parties", "🎯 The Work (Scope)", "💰 The Money"])

with tab1:
    c1, c2 = st.columns(2)
    with c1:
        freelancer_name = st.text_input("Provider Name (You)", "Amit Kumar", help="Your full legal name as it appears on your bank account")
        cities = [
            "Bengaluru, Karnataka", 
            "New Delhi, Delhi", 
            "Mumbai, Maharashtra", 
            "Chennai, Tamil Nadu", 
            "Hyderabad, Telangana", 
            "Pune, Maharashtra", 
            "Kolkata, West Bengal",
            "Ahmedabad, Gujarat",
            "Jaipur, Rajasthan",
            "Other (Type Manually)"
        ]
        selected_city = st.selectbox("Your City (Jurisdiction)", cities, help="This determines which court has jurisdiction if disputes arise")
        jurisdiction_city = st.text_input("Type Your City", "Mysuru, Karnataka") if selected_city == "Other (Type Manually)" else selected_city
    with c2:
        client_name = st.text_input("Client Name", "Tech Solutions Pvt Ltd", help="Company name or individual's full name")
        gst_registered = st.checkbox("I am GST Registered", help="Check this if you have a valid GSTIN number")

with tab2:
    st.markdown('<div class="warning-box">⚠️ <b>IMPORTANT:</b> Selecting a category automatically adjusts the <b>Legal Clauses</b> (IP Rights, Warranty, Cancellation Terms) to match your industry risks. For example, Tech contracts get bug fix warranties, Creative contracts get source file protection.</div>', unsafe_allow_html=True)
    
    template_choice = st.selectbox(
        "✨ Select Your Industry (This Activates Smart Clauses):", 
        list(scope_templates.keys()), 
        key="template_selector", 
        on_change=update_scope, 
        help="The contract automatically adapts to your industry's legal risks"
    )
    
    scope_work = st.text_area(
        "Scope of Work (Annexure A)", 
        key="scope_text", 
        height=220, 
        help="Be SPECIFIC. Vague contracts lead to unpaid work. Include: deliverables, formats, timelines, revisions, and exclusions."
    )
    
    if template_choice != "Select a template...":
        st.caption("💡 **Pro Tip:** Edit the template above to match your exact project. The more specific you are, the better protected you'll be.")

with tab3:
    c1, c2, c3 = st.columns(3)
    with c1: 
        project_fee_num = st.number_input(
            "Total Project Fee (INR)", 
            value=50000, 
            step=1000, 
            help="Total contract value (before GST)"
        )
    with c2: 
        hourly_rate_num = st.number_input(
            "Overtime Rate (INR/hr)", 
            value=2000, 
            step=500, 
            help="Rate for scope creep and out-of-scope requests"
        )
    with c3:
        advance_percent = st.slider(
            "Advance Required (%)", 
            0, 100, 50, 
            help="Industry standard: 30-50% for freelancers. Never start work without advance payment."
        )
    
    advance_amount = int(project_fee_num * (advance_percent/100))
    balance_amount = project_fee_num - advance_amount
    
    st.markdown(f"""
    <div style="background: rgba(59, 130, 246, 0.1); padding: 15px; border-radius: 8px; border-left: 4px solid #3b82f6;">
        <b>💰 Payment Breakdown:</b><br/>
        • Advance Payment (Due Before Work Starts): <b>Rs. {advance_amount:,}</b><br/>
        • Balance Payment (Due on Delivery): <b>Rs. {balance_amount:,}</b><br/>
        • Total Project Value: <b>Rs. {project_fee_num:,}</b>
    </div>
    """, unsafe_allow_html=True)

st.markdown("---")

# --- CONSENT & GENERATION ---
check_terms = st.checkbox("✅ I agree to the Terms of Use & Privacy Policy. I understand this is an automated tool, not legal advice from a lawyer.")

c_main = st.columns([1, 2, 1])
with c_main[1]: 
    if check_terms:
        generate_btn = st.button("🚀 Generate Legal Contract Now", type="primary")
    else:
        st.button("🚀 Generate Legal Contract Now", disabled=True, help="Please accept the Terms to proceed.")
        generate_btn = False

if generate_btn:
    # VALIDATION: Check if template is selected
    if template_choice == "Select a template...":
        st.error("⚠️ **Please select an Industry Template** in the 'The Work (Scope)' tab before generating.")
        st.stop()
    
    # VALIDATION: Check if scope is filled
    if not st.session_state.scope_text.strip():
        st.error("⚠️ **Scope of Work is empty!** Please fill in the deliverables in the 'The Work' tab.")
        st.stop()
    
    with st.spinner("🔐 Drafting your watertight contract..."):
        time.sleep(1.5)
    
    # Format currency safely
    safe_cost = f"Rs. {project_fee_num:,}"
    safe_rate = f"Rs. {hourly_rate_num:,}"
    safe_scope = st.session_state.scope_text.replace("₹", "Rs. ")
    gst_clause = "(Exclusive of GST)" if gst_registered else ""
    
    # Get smart clauses
    smart = get_smart_clauses(template_choice, safe_rate)
    
    # Build contract text with enhanced formatting
    full_text = f"""PROFESSIONAL SERVICE AGREEMENT

Date: {datetime.date.today().strftime('%B %d, %Y')}

BETWEEN: {freelancer_name} ("Provider")
AND: {client_name} ("Client")

This Agreement governs the professional services to be provided by the Provider to the Client as detailed in Annexure A.

═══════════════════════════════════════════════════════════════

1. PAYMENT TERMS & LATE PAYMENT INTEREST (MSME ACT COMPLIANCE)

Total Project Fee: {safe_cost} {gst_clause}
Advance Payment Required: {advance_percent}% (Rs. {advance_amount:,})
Balance Payment: Rs. {balance_amount:,} (due upon delivery of final deliverables)

CRITICAL PAYMENT PROTECTION: As per Section 16 of the Micro, Small and Medium Enterprises Development (MSMED) Act, 2006, any payment delayed beyond the agreed date will automatically attract compound interest at THREE TIMES (3x) the Bank Rate notified by the Reserve Bank of India. This is non-negotiable and legally enforceable.

Payment must be made via bank transfer, UPI, or other traceable methods. Cash payments are not accepted.

═══════════════════════════════════════════════════════════════

2. ACCEPTANCE & REVISIONS

{smart['acceptance']}

Revisions must be requested in writing (email/WhatsApp) with specific, actionable feedback. Vague feedback such as "make it better" or "I don't like it" without concrete direction will not be considered valid revision requests.

═══════════════════════════════════════════════════════════════

3. INTELLECTUAL PROPERTY RIGHTS (IP LOCK)

{smart['ip_rights']}

CRITICAL: Until full and final payment is received, all work product, drafts, source files, and deliverables remain the exclusive property of the Provider. Any use, publication, or distribution of work before payment completion constitutes Copyright Infringement under the Copyright Act, 1957, and Provider reserves the right to pursue legal remedies.

═══════════════════════════════════════════════════════════════

4. WARRANTY & SUPPORT

{smart['warranty']}

═══════════════════════════════════════════════════════════════

5. CONFIDENTIALITY (NON-DISCLOSURE AGREEMENT)

Both parties agree to maintain strict confidentiality regarding all proprietary information, trade secrets, business strategies, client lists, source code, and sensitive data shared during the course of this engagement.

This obligation survives for TWO (2) YEARS after termination of this Agreement.

Exceptions: Does not apply to information that is (a) publicly available, (b) already known to the receiving party, or (c) required to be disclosed by law.

═══════════════════════════════════════════════════════════════

6. COMMUNICATION POLICY & ANTI-GHOSTING CLAUSE

Provider Response Time: Provider commits to responding to Client communications within 1 (one) business day (Monday-Friday, excluding national holidays).

Client Responsiveness: Client agrees to provide timely feedback, approvals, and required materials.

ANTI-GHOSTING PROTECTION: If Client fails to respond to Provider communications for FOURTEEN (14) consecutive days without prior notice, the project will be deemed TERMINATED. All payments made are NON-REFUNDABLE. Provider reserves the right to charge a standby fee of Rs. {int(hourly_rate_num * 8):,} per day for extended delays caused by Client unavailability.

═══════════════════════════════════════════════════════════════

7. CANCELLATION & KILL FEE

{smart['cancellation']}

═══════════════════════════════════════════════════════════════

8. TERMINATION BY PROVIDER

{smart.get('termination', 'Provider may terminate with 7 days written notice if Client breaches payment terms. Client owes payment for work completed till termination date.')}

═══════════════════════════════════════════════════════════════

9. FORCE MAJEURE

Neither party shall be held liable for failure or delay in performance due to circumstances beyond reasonable control including but not limited to:
- Acts of God (earthquakes, floods, pandemics)
- Internet/telecommunications/power failures
- Government actions, war, or civil unrest
- Labor strikes or supply chain disruptions

Upon occurrence of such events, the affected party must notify the other within 48 hours.

═══════════════════════════════════════════════════════════════

10. LIMITATION OF LIABILITY

Provider's total liability under this Agreement is strictly limited to the Total Fee paid by Client. Provider shall NOT be liable for any indirect, incidental, consequential, or punitive damages including lost profits, business interruption, or data loss.

═══════════════════════════════════════════════════════════════

11. DISPUTE RESOLUTION & JURISDICTION

Governing Law: This Agreement is governed by the laws of India.

Arbitration: Any disputes arising from this Agreement shall first be attempted to be resolved through good-faith negotiation within 15 days. If unresolved, disputes shall be referred to BINDING ARBITRATION under the Arbitration and Conciliation Act, 1996.

Arbitration Seat: {jurisdiction_city}
Language: English
Single Arbitrator: Mutually appointed by both parties

The prevailing party in arbitration shall be entitled to recover reasonable legal costs and arbitration fees.

═══════════════════════════════════════════════════════════════

12. GST COMPLIANCE

All fees quoted are EXCLUSIVE of Goods and Services Tax (GST). Current applicable GST rates will be added to invoices as per Indian tax laws.

Client Responsibility: Client is responsible for providing valid GSTIN details if claiming Input Tax Credit. Any penalties arising from incorrect GST information provided by Client shall be Client's responsibility.

═══════════════════════════════════════════════════════════════

13. STAMP DUTY COMPLIANCE

NOTE: This agreement may require stamp duty payment as per the Indian Stamp Act, 1899. Stamp duty rates vary by state. Client is responsible for ensuring stamp duty compliance in their state of residence.

═══════════════════════════════════════════════════════════════

14. ENTIRE AGREEMENT

This Agreement (including Annexure A) constitutes the entire understanding between the parties and supersedes all prior discussions, agreements, or understandings. Any modifications must be made in writing and signed by both parties.

═══════════════════════════════════════════════════════════════

SIGNATURES

By signing below, both parties acknowledge that they have read, understood, and agree to be bound by all terms and conditions of this Agreement.

PROVIDER:
Signature: _____________________
Name: {freelancer_name}
Date: _____________________

CLIENT:
Signature: _____________________
Name: {client_name}
Date: _____________________
"""

    # Generate professional documents
    try:
        pdf_data = create_professional_pdf(full_text, safe_scope, freelancer_name, client_name)
        docx_data = create_professional_docx(full_text, safe_scope, freelancer_name, client_name)
        
        # SUCCESS UI
        st.balloons()
        st.success("✅ **Contract Generated Successfully!** Your legal protection is ready.")
        
        st.markdown("""
        <div class="success-box">
            <b>📥 Next Steps:</b><br/>
            1. Download both PDF (for signing) and Word (for editing if needed)<br/>
            2. Send to your client via email or WhatsApp<br/>
            3. Get both signatures (digital signatures are legally valid in India)<br/>
            4. Keep a copy for your records<br/>
            5. <b>DO NOT START WORK until you receive the advance payment!</b>
        </div>
        """, unsafe_allow_html=True)
        
        col_d1, col_d2 = st.columns(2)
        with col_d1:
            st.download_button(
                "📄 Download as PDF (Recommended for Signing)", 
                data=pdf_data, 
                file_name=f"Contract_{client_name.replace(' ', '_')}_{datetime.date.today().strftime('%Y%m%d')}.pdf", 
                mime="application/pdf", 
                use_container_width=True
            )
        with col_d2:
            st.download_button(
                "📝 Download as Word (Editable)", 
                data=docx_data, 
                file_name=f"Contract_{client_name.replace(' ', '_')}_{datetime.date.today().strftime('%Y%m%d')}.docx", 
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                use_container_width=True
            )
        
        st.markdown("---")
        
        # Preview
        with st.expander("👀 Preview Full Contract (Click to Expand)"):
            st.text_area("", value=full_text + "\n\n" + "═"*60 + "\nANNEXURE A: SCOPE OF WORK\n" + "═"*60 + "\n\n" + safe_scope, height=400, disabled=True)
        
        # Share options
        st.markdown("### 📤 Share with Client")
        col_share1, col_share2 = st.columns(2)
        with col_share1:
            whatsapp_text = f"Hi {client_name}, please find our service agreement attached. Let me know if you have any questions!"
            whatsapp_link = f"https://wa.me/?text={whatsapp_text}"
            st.link_button("📱 Share on WhatsApp", whatsapp_link, use_container_width=True)
        with col_share2:
            email_subject = f"Service Agreement - {freelancer_name} & {client_name}"
            email_body = f"Dear {client_name},%0D%0A%0D%0APlease find attached our Professional Service Agreement for your review.%0D%0A%0D%0AKey Terms:%0D%0A- Total Fee: {safe_cost}%0D%0A- Advance: {advance_percent}% (Rs. {advance_amount:,})%0D%0A- Timeline: As per Scope of Work%0D%0A%0D%0APlease review and sign. Work will commence upon receipt of advance payment.%0D%0A%0D%0ABest regards,%0D%0A{freelancer_name}"
            mailto_link = f"mailto:?subject={email_subject}&body={email_body}"
            st.link_button("📧 Send via Email", mailto_link, use_container_width=True)
    
    except Exception as e:
        st.error(f"❌ Error generating contract: {str(e)}")
        st.error("Please try again or contact support if the issue persists.")

# --- LEGAL FOOTER ---
st.markdown("---")
st.markdown("""
<div class="legal-footer">
    <p><b>© 2025 Freelance Shield Pro.</b> All rights reserved. Built with ❤️ for Indian Freelancers.</p>
    <p><b>Legal Disclaimer:</b> This tool provides automated contract templates for informational purposes only and does not constitute legal advice. 
    We are not a law firm and do not create attorney-client relationships. For projects exceeding Rs. 1,00,000 or complex legal situations, 
    please consult a qualified attorney licensed to practice in India. By using this tool, you acknowledge that you understand its limitations.</p>
    <p style="margin-top: 10px; font-size: 0.75rem; color: #475569;">
        Made by a law student passionate about freelancer rights | 
        <a href="https://wa.me/YOUR_NUMBER_HERE" style="color: #3b82f6;">Contact for Custom Contracts</a> | 
        <a href="#" style="color: #3b82f6;">Report a Bug</a>
    </p>
</div>
""", unsafe_allow_html=True)